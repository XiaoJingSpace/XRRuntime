#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 转 Word 文档转换脚本
将所有章节合并为一个标准技术书籍格式的 .docx 文件
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import hashlib

# 章节顺序
CHAPTER_ORDER = [
    # 第一部分：基础篇
    ('book/chapters/part1/chapter01.md', '第1章：OpenXR Runtime 概述'),
    ('book/chapters/part1/chapter02.md', '第2章：开发环境准备'),
    
    # 第二部分：编译构建篇
    ('book/chapters/part2/chapter03.md', '第3章：项目结构与构建系统'),
    ('book/chapters/part2/chapter04.md', '第4章：编译实践'),
    ('book/chapters/part2/chapter05.md', '第5章：输出产物与验证'),
    
    # 第三部分：功能实现篇
    ('book/chapters/part3/chapter06.md', '第6章：XRRuntime 框架架构设计'),
    ('book/chapters/part3/chapter07.md', '第7章：Instance 管理实现'),
    ('book/chapters/part3/chapter08.md', '第8章：Session 管理实现'),
    ('book/chapters/part3/chapter09.md', '第9章：Frame 循环实现'),
    ('book/chapters/part3/chapter10.md', '第10章：追踪系统实现'),
    ('book/chapters/part3/chapter11.md', '第11章：渲染系统实现'),
    ('book/chapters/part3/chapter12.md', '第12章：输入系统实现'),
    ('book/chapters/part3/chapter13.md', '第13章：事件系统实现'),
    ('book/chapters/part3/chapter14.md', '第14章：平台集成实现'),
    
    # 第四部分：输出引用篇
    ('book/chapters/part4/chapter15.md', '第15章：运行时库文件部署'),
    ('book/chapters/part4/chapter16.md', '第16章：OpenXR Loader 配置'),
    ('book/chapters/part4/chapter17.md', '第17章：Java/Kotlin 集成'),
    ('book/chapters/part4/chapter18.md', '第18章：OpenXR API 调用实践'),
    ('book/chapters/part4/chapter19.md', '第19章：性能优化与调试'),
    
    # 第五部分：高级主题篇
    ('book/chapters/part5/chapter20.md', '第20章：框架扩展设计'),
    ('book/chapters/part5/chapter21.md', '第21章：测试与验证'),
    ('book/chapters/part5/chapter22.md', '第22章：故障排除与维护'),
    ('book/chapters/part5/chapter23.md', '第23章：最佳实践与总结'),
    
    # 附录
    ('book/appendices/appendix_a.md', '附录A：OpenXR API 快速参考'),
    ('book/appendices/appendix_b.md', '附录B：QVR API 参考'),
    ('book/appendices/appendix_c.md', '附录C：CMake 变量参考'),
    ('book/appendices/appendix_d.md', '附录D：Gradle 配置参考'),
    ('book/appendices/appendix_e.md', '附录E：故障排除速查表'),
    ('book/appendices/appendix_f.md', '附录F：资源链接'),
]

def setup_document_styles(doc):
    """设置文档样式"""
    styles = doc.styles
    
    # 正文样式
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '宋体'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(6)
    
    # 列表样式优化
    try:
        list_bullet = styles['List Bullet']
        list_bullet.font.name = '宋体'
        list_bullet.font.size = Pt(12)
        list_number = styles['List Number']
        list_number.font.name = '宋体'
        list_number.font.size = Pt(12)
    except:
        pass

def parse_markdown(content):
    """解析 Markdown 内容"""
    lines = content.split('\n')
    elements = []
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 标题1 (#)
        if line.startswith('# ') and not line.startswith('##'):
            elements.append(('h1', line[2:].strip()))
            i += 1
        # 标题2 (##)
        elif line.startswith('## ') and not line.startswith('###'):
            elements.append(('h2', line[3:].strip()))
            i += 1
        # 标题3 (###)
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
            i += 1
        # 代码块
        elif line.startswith('```'):
            code_lines = []
            lang = line[3:].strip() if len(line) > 3 else ''
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            # 检查是否是 Mermaid 图表
            if lang == 'mermaid':
                elements.append(('mermaid', '\n'.join(code_lines)))
            else:
                elements.append(('code', '\n'.join(code_lines)))
        # 有序列表项
        elif re.match(r'^\d+\.\s+', line):
            items = []
            items.append(re.sub(r'^\d+\.\s+', '', line))
            i += 1
            # 收集连续列表项
            while i < len(lines):
                next_line = lines[i].rstrip()
                if re.match(r'^\d+\.\s+', next_line):
                    items.append(re.sub(r'^\d+\.\s+', '', next_line))
                    i += 1
                elif next_line.startswith('- ') or next_line.startswith('* '):
                    break
                elif next_line.startswith('#') or next_line.startswith('```'):
                    break
                elif not next_line.strip():
                    i += 1
                    break
                else:
                    # 续行
                    items[-1] += ' ' + next_line.strip()
                    i += 1
            elements.append(('olist', items))
        # 无序列表项
        elif line.startswith('- ') or line.startswith('* '):
            items = []
            items.append(line[2:].strip())
            i += 1
            # 收集连续列表项
            while i < len(lines):
                next_line = lines[i].rstrip()
                if next_line.startswith('- ') or next_line.startswith('* '):
                    items.append(next_line[2:].strip())
                    i += 1
                elif re.match(r'^\d+\.\s+', next_line):
                    break
                elif next_line.startswith('#') or next_line.startswith('```'):
                    break
                elif not next_line.strip():
                    i += 1
                    break
                else:
                    # 续行
                    items[-1] += ' ' + next_line.strip()
                    i += 1
            elements.append(('list', items))
        # 空行
        elif not line.strip():
            elements.append(('blank', ''))
            i += 1
        # 普通段落
        else:
            para_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and \
                  not lines[i].startswith('- ') and not lines[i].startswith('* ') and \
                  not lines[i].startswith('```') and not re.match(r'^\d+\.\s+', lines[i]):
                para_lines.append(lines[i].rstrip())
                i += 1
            if para_lines:
                elements.append(('para', ' '.join(para_lines)))
            else:
                i += 1
    
    return elements

def _add_text_with_formatting(p, text):
    """添加带格式的文本（处理加粗和行内代码）"""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 加粗文本
            run = p.add_run(part[2:-2])
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # 行内代码
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.bold = False
        elif part:
            # 普通文本
            run = p.add_run(part)
            run.font.name = '宋体'
            run.font.size = Pt(12)

def _add_mermaid_placeholder(doc, content):
    """添加 Mermaid 图表占位符"""
    p = doc.add_paragraph()
    run = p.add_run('[Mermaid 图表 - 需要转换为图片]')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(255, 128, 0)  # 橙色提示
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    # 添加说明
    note_p = doc.add_paragraph()
    note_run = note_p.add_run('注：请使用 mermaid_to_image.py 脚本将图表转换为图片，或手动在 Mermaid Live Editor 中转换后插入。')
    note_run.font.name = '宋体'
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(128, 128, 128)

def add_elements_to_doc(doc, elements):
    """将解析的元素添加到文档"""
    for elem_type, content in elements:
        if elem_type == 'h1':
            # 章节标题：居中，前后分页
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.font.name = '黑体'
            run.font.size = Pt(20)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(18)
            # 章节标题后添加分页符（除了第一个）
            if len([e for e in elements if e[0] == 'h1']) > 1:
                p = doc.add_paragraph()
                p.paragraph_format.page_break_before = True
        elif elem_type == 'h2':
            # 节标题：左对齐，加粗
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = '黑体'
            run.font.size = Pt(16)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Cm(0)
        elif elem_type == 'h3':
            # 小节标题：左对齐，加粗
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = '黑体'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0)
        elif elem_type == 'code':
            # 代码块：等宽字体，灰色背景，缩进
            p = doc.add_paragraph()
            run = p.add_run(content.rstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.right_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0)
            # 设置代码块背景色和边框
            try:
                shading = p._element.get_or_add_pPr().get_or_add_shd()
                shading.set(qn('w:fill'), 'F0F0F0')
            except:
                pass
        elif elem_type == 'mermaid':
            # Mermaid 图表：尝试插入图片
            # 生成图表文件名（基于内容哈希）
            diagram_hash = hashlib.md5(content.encode('utf-8')).hexdigest()[:8]
            diagram_name = f"mermaid_{diagram_hash}.png"
            diagram_path = Path('book/images/mermaid') / diagram_name
            
            # 检查图片是否存在
            if diagram_path.exists():
                try:
                    # 插入图片
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(diagram_path), width=Cm(14))  # 宽度14cm
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                except Exception as e:
                    print(f"插入图片失败: {e}")
                    # 如果插入失败，显示占位符
                    _add_mermaid_placeholder(doc, content)
            else:
                # 图片不存在，显示占位符和说明
                _add_mermaid_placeholder(doc, content)
                print(f"提示: Mermaid 图表图片不存在: {diagram_path}")
                print(f"  请运行: python scripts/mermaid_to_image.py <markdown_file>")
        elif elem_type == 'list':
            # 无序列表：项目符号
            for idx, item in enumerate(content):
                p = doc.add_paragraph(style='List Bullet')
                # 去掉 ** 标记，直接显示文本
                clean_item = re.sub(r'\*\*([^*]+)\*\*', r'\1', item)
                run = p.add_run(clean_item)
                run.font.name = '宋体'
                run.font.size = Pt(12)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.space_after = Pt(3)
        elif elem_type == 'olist':
            # 有序列表：数字编号
            for idx, item in enumerate(content):
                p = doc.add_paragraph(style='List Number')
                # 去掉 ** 标记，直接显示文本
                clean_item = re.sub(r'\*\*([^*]+)\*\*', r'\1', item)
                run = p.add_run(clean_item)
                run.font.name = '宋体'
                run.font.size = Pt(12)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.space_after = Pt(3)
        elif elem_type == 'para':
            # 普通段落：处理行内代码，去掉加粗标记
            para_text = content
            p = doc.add_paragraph()
            
            # 去掉 ** 标记，处理行内代码
            # 先处理行内代码（保留），再去掉加粗标记
            parts = re.split(r'(`[^`]+`)', para_text)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    # 行内代码
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif part:
                    # 普通文本，去掉 ** 标记
                    clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', part)
                    run = p.add_run(clean_text)
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
        elif elem_type == 'blank':
            # 空行：减少间距
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

def main():
    """主函数"""
    # 创建文档
    doc = Document()
    
    # 设置页面
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # 设置样式
    setup_document_styles(doc)
    
    # 添加封面页
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('XRRuntime框架设计与开发实践指南')
    title_run.font.name = '黑体'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_para.paragraph_format.space_before = Pt(250)
    
    subtitle_para = doc.add_paragraph('完整版')
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.runs[0]
    subtitle_run.font.name = '宋体'
    subtitle_run.font.size = Pt(16)
    subtitle_para.paragraph_format.space_after = Pt(350)
    
    # 添加分页符
    doc.add_page_break()
    
    # 添加目录页（占位）
    toc_para = doc.add_paragraph('目录', style='Heading 1')
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_para.paragraph_format.space_after = Pt(300)
    doc.add_page_break()
    
    # 处理所有章节
    for file_path, chapter_title in CHAPTER_ORDER:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            print(f"警告: 文件不存在: {file_path}")
            continue
        
        print(f"处理: {chapter_title}")
        
        # 读取文件内容
        with open(file_path_obj, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 解析 Markdown
        elements = parse_markdown(content)
        
        # 添加到文档
        add_elements_to_doc(doc, elements)
        
        # 章节之间添加分页符（除了最后一个和附录）
        if chapter_title != CHAPTER_ORDER[-1][1]:
            # 检查是否是附录开始
            if chapter_title.startswith('附录'):
                # 附录前添加分页符
                p = doc.add_paragraph()
                p.paragraph_format.page_break_before = True
            elif not chapter_title.startswith('附录'):
                # 普通章节之间添加分页符
                doc.add_page_break()
    
    # 保存文档
    output_path = Path('book/XRRuntime框架设计与开发实践指南.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # 如果文件存在且被占用，先保存为临时文件
    temp_path = Path('book/XRRuntime框架设计与开发实践指南_temp.docx')
    try:
        doc.save(str(output_path))
        print(f"\n文档已保存: {output_path}")
    except PermissionError:
        print("原文件被占用，保存为临时文件...")
        doc.save(str(temp_path))
        print(f"文档已保存为临时文件: {temp_path}")
        print("请关闭原文件后，将临时文件重命名为正式文件名")

if __name__ == '__main__':
    main()

